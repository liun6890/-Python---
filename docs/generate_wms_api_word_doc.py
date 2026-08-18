from __future__ import annotations

import json
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
HTML_PATH = ROOT / "wms-api-tester.html"
OUTPUT_PATH = ROOT / "WMS系统完整接口文档.docx"
BASE_URL = "http://127.0.0.1:8000"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_document_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_table(document: Document, rows: list[list[str]], header_rows: set[int] | None = None):
    header_rows = header_rows or set()
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_text(cell, value, bold=row_idx in header_rows, center=row_idx in header_rows)
            if row_idx in header_rows:
                set_cell_shading(cell, "D9D9D9")
    document.add_paragraph("")
    return table


def extract_default_collection() -> list[dict]:
    html = HTML_PATH.read_text(encoding="utf-8")
    match = re.search(r"const defaultCollection = (\[[\s\S]*?\]);\s*let collection", html)
    if not match:
        raise RuntimeError("未找到 wms-api-tester.html 中的 defaultCollection")

    js_code = f"""
const defaultCollection = {match.group(1)};
console.log(JSON.stringify(defaultCollection));
"""
    with tempfile.NamedTemporaryFile("w", suffix=".js", delete=False, encoding="utf-8") as temp:
        temp.write(js_code)
        temp_path = Path(temp.name)

    try:
        result = subprocess.run(
            ["node", str(temp_path)],
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
        )
        return json.loads(result.stdout)
    finally:
        temp_path.unlink(missing_ok=True)


def flatten_json_fields(value, prefix: str = "") -> list[tuple[str, object]]:
    fields: list[tuple[str, object]] = []
    if isinstance(value, dict):
        for key, child in value.items():
            name = f"{prefix}.{key}" if prefix else key
            if isinstance(child, (dict, list)):
                fields.extend(flatten_json_fields(child, name))
            else:
                fields.append((name, child))
    elif isinstance(value, list):
        item_prefix = f"{prefix}[]" if prefix else "[]"
        if value:
            fields.extend(flatten_json_fields(value[0], item_prefix))
        else:
            fields.append((item_prefix, "array"))
    return fields


def value_type(value) -> str:
    if isinstance(value, bool):
        return "boolean"
    if isinstance(value, int):
        return "number"
    if isinstance(value, float):
        return "number"
    if isinstance(value, list):
        return "array"
    if isinstance(value, dict):
        return "object"
    if value is None:
        return "null"
    return "string"


def build_input_rows(item: dict) -> list[list[str]]:
    rows = [["中文名", "英文名", "参数类型及取值范围", "是否必填", "说明"]]
    if item.get("auth"):
        rows.append(["访问令牌", "Authorization", "Header: Bearer <token>", "是", "登录接口返回的 JWT Token"])

    for param in item.get("params") or []:
        rows.append([
            param.get("desc") or param.get("key") or "查询参数",
            param.get("key", ""),
            f"Query: {param.get('type', 'string')}",
            "是" if param.get("enabled") else "否",
            param.get("desc", ""),
        ])

    raw_body = (item.get("body") or "").strip()
    if raw_body:
        try:
            body = json.loads(raw_body)
            flattened = flatten_json_fields(body)
            for name, sample in flattened:
                rows.append([
                    name.split(".")[-1].replace("[]", "明细"),
                    name,
                    f"Body: {value_type(sample)}",
                    "是",
                    f"示例值：{sample}",
                ])
        except json.JSONDecodeError:
            rows.append(["请求体", "body", "Body: JSON", "是", raw_body])

    if len(rows) == 1:
        rows.append(["无", "-", "-", "否", "该接口无显式输入参数"])
    return rows


def build_output_rows(item: dict) -> list[list[str]]:
    rows = [
        ["中文名", "英文名", "参数类型及取值范围", "是否必填", "说明"],
        ["业务返回码", "code", "number", "是", "200 表示业务成功，400/401/500 表示失败"],
        ["返回消息", "message", "string", "是", "success 或错误消息"],
        ["数据内容", "data", "object/array/null", "是", "接口业务数据；等同模板中的 datas"],
    ]

    seen = {"code", "message", "data"}
    for assertion in item.get("assertions") or []:
        path = assertion.get("path") or ""
        if not path or path in seen or path == "code":
            continue
        seen.add(path)
        if path.startswith("data."):
            cn_name = path.replace("data.", "数据.")
        else:
            cn_name = path
        op = assertion.get("operator", "")
        type_hint = "array" if op == "isArray" else "object/string/number"
        rows.append([cn_name, path, type_hint, "否", "自动化断言关注字段"])
    return rows


def request_kind(item: dict) -> str:
    if item.get("params"):
        return "Query 参数"
    if (item.get("body") or "").strip():
        return "JSON Body"
    return "无"


def add_api_section(document: Document, index: int, item: dict) -> None:
    add_heading(document, f"{index}. {item['name']}", level=2)
    path = item["path"]
    method = item.get("method", "")
    full_url = f"{BASE_URL}{path}"
    base_rows = [
        ["名称", item["name"]],
        ["说明", item.get("description", "")],
        ["调用方系统", "WMS 前端系统 / 自动化测试脚本"],
        ["请求方式", method],
        ["调用路径", full_url],
        ["认证方式", "Bearer Token" if item.get("auth") else "无需认证"],
        ["调用参数", request_kind(item)],
        ["返回参数", "code / message / data"],
    ]
    add_table(document, base_rows, header_rows=set())

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("输入参数")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    add_table(document, build_input_rows(item), header_rows={0})

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("输出参数（返回结果）")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    add_table(document, build_output_rows(item), header_rows={0})

    if (item.get("body") or "").strip():
        add_heading(document, "请求体示例", level=3)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(item["body"])
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    if item.get("notes"):
        add_heading(document, "业务说明", level=3)
        for note in item["notes"]:
            document.add_paragraph(str(note), style="List Bullet")


def build_document(collection: list[dict]) -> Document:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    set_document_font(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("WMS 系统完整接口文档")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_paragraph("文档来源：根据当前 WMS 后端路由和独立接口测试工作台生成。")
    document.add_paragraph(f"基础地址：{BASE_URL}/api")
    document.add_paragraph("认证方式：除登录接口外，业务接口默认使用 Authorization: Bearer <token>。")
    document.add_paragraph("通用返回：code/message/data，其中 data 对应截图模板中的 datas 数据内容。")

    add_heading(document, "接口总览", level=1)
    overview_rows = [["序号", "模块", "接口名称", "请求方式", "路径", "认证"]]
    for idx, item in enumerate(collection, start=1):
        overview_rows.append([
            str(idx),
            item.get("module", ""),
            item.get("name", ""),
            item.get("method", ""),
            item.get("path", ""),
            "是" if item.get("auth") else "否",
        ])
    add_table(document, overview_rows, header_rows={0})

    add_heading(document, "接口明细", level=1)
    for idx, item in enumerate(collection, start=1):
        add_api_section(document, idx, item)

    return document


def main() -> None:
    collection = extract_default_collection()
    document = build_document(collection)
    document.save(OUTPUT_PATH)
    print(f"generated: {OUTPUT_PATH}")
    print(f"api_count: {len(collection)}")


if __name__ == "__main__":
    main()
